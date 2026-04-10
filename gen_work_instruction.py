"""
Generate Professional 3D Camera & Navigation Controls – Work Instruction Document
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Colour palette ──────────────────────────────────────────────────────────
DARK_BLUE   = RGBColor(0x1F, 0x3B, 0x6E)   # heading 1
MID_BLUE    = RGBColor(0x2E, 0x75, 0xB6)   # heading 2
LIGHT_BLUE  = RGBColor(0xD5, 0xE8, 0xF5)   # table header fill
PALE_BLUE   = RGBColor(0xEE, 0xF4, 0xFB)   # alternate row
ORANGE      = RGBColor(0xC5, 0x5A, 0x11)   # heading 3 accent
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
BLACK       = RGBColor(0x00, 0x00, 0x00)
DARK_GREY   = RGBColor(0x40, 0x40, 0x40)

def hex_to_rgb_str(rgb: RGBColor) -> str:
    return f"{rgb.r:02X}{rgb.g:02X}{rgb.b:02X}"

# ── Low-level helpers ────────────────────────────────────────────────────────

def set_cell_bg(cell, rgb: RGBColor):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_to_rgb_str(rgb))
    tcPr.append(shd)

def set_cell_borders(cell, color="CCCCCC", sz=4):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top','left','bottom','right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),   'single')
        el.set(qn('w:sz'),    str(sz))
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)

def set_cell_margins(cell, top=80, bottom=80, left=120, right=120):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    mar  = OxmlElement('w:tcMar')
    for side, val in (('top',top),('left',left),('bottom',bottom),('right',right)):
        m = OxmlElement(f'w:{side}')
        m.set(qn('w:w'),    str(val))
        m.set(qn('w:type'), 'dxa')
        mar.append(m)
    tcPr.append(mar)

def set_col_width(cell, width_twips: int):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW  = OxmlElement('w:tcW')
    tcW.set(qn('w:w'),    str(width_twips))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)

def bold_run(para, text, color=None, size=None):
    run = para.add_run(text)
    run.bold = True
    if color: run.font.color.rgb = color
    if size:  run.font.size = Pt(size)
    return run

def add_para_bottom_border(para, color="2E75B6", size=6):
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    str(size))
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), color)
    pBdr.append(bot)
    pPr.append(pBdr)

def set_spacing(para, before=0, after=0, line=None):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after  = Pt(after)
    if line:
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(line)

# ── Document-level helpers ───────────────────────────────────────────────────

def h1(doc, text):
    p = doc.add_paragraph()
    set_spacing(p, before=18, after=4)
    add_para_bottom_border(p, color=hex_to_rgb_str(DARK_BLUE), size=8)
    run = p.add_run(text)
    run.bold  = True
    run.font.size  = Pt(16)
    run.font.color.rgb = DARK_BLUE
    run.font.name = "Arial"
    return p

def h2(doc, text):
    p = doc.add_paragraph()
    set_spacing(p, before=10, after=3)
    run = p.add_run(text)
    run.bold  = True
    run.font.size  = Pt(13)
    run.font.color.rgb = MID_BLUE
    run.font.name = "Arial"
    return p

def h3(doc, text):
    p = doc.add_paragraph()
    set_spacing(p, before=8, after=2)
    run = p.add_run(text)
    run.bold  = True
    run.font.size  = Pt(11)
    run.font.color.rgb = ORANGE
    run.font.name = "Arial"
    return p

def body(doc, text, indent=False):
    p = doc.add_paragraph()
    set_spacing(p, before=2, after=2)
    run = p.add_run(text)
    run.font.size  = Pt(10)
    run.font.color.rgb = DARK_GREY
    run.font.name = "Arial"
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    return p

def bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    set_spacing(p, before=1, after=1)
    p.paragraph_format.left_indent  = Cm(0.5 + level * 0.5)
    p.paragraph_format.first_line_indent = Cm(-0.35)
    run = p.add_run(text)
    run.font.size  = Pt(10)
    run.font.color.rgb = DARK_GREY
    run.font.name = "Arial"
    return p

def note(doc, text):
    p = doc.add_paragraph()
    set_spacing(p, before=3, after=3)
    p.paragraph_format.left_indent = Cm(0.4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    lft  = OxmlElement('w:left')
    lft.set(qn('w:val'),   'single')
    lft.set(qn('w:sz'),    '12')
    lft.set(qn('w:space'), '4')
    lft.set(qn('w:color'), hex_to_rgb_str(MID_BLUE))
    pBdr.append(lft)
    pPr.append(pBdr)
    run = p.add_run("NOTE  ")
    run.bold = True
    run.font.color.rgb = MID_BLUE
    run.font.size = Pt(9)
    run.font.name = "Arial"
    run2 = p.add_run(text)
    run2.font.size = Pt(9)
    run2.font.color.rgb = DARK_GREY
    run2.font.name = "Arial"
    return p

def settings_table(doc, rows_data, col_widths=None):
    """rows_data = list of (setting, type, default, description)"""
    if col_widths is None:
        col_widths = [2400, 1400, 1400, 4000]  # twips
    headers = ["Setting / Property", "Type", "Default", "Description"]
    table = doc.add_table(rows=1 + len(rows_data), cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = 'Table Grid'

    # Header row
    hrow = table.rows[0]
    for i, (cell, hdr, w) in enumerate(zip(hrow.cells, headers, col_widths)):
        set_cell_bg(cell, DARK_BLUE)
        set_cell_borders(cell, color="FFFFFF", sz=4)
        set_cell_margins(cell)
        set_col_width(cell, w)
        p = cell.paragraphs[0]
        run = p.add_run(hdr)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(9)
        run.font.name = "Arial"
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Data rows
    for ri, row_data in enumerate(rows_data):
        row = table.rows[ri + 1]
        fill = PALE_BLUE if ri % 2 == 0 else WHITE
        for ci, (cell, val, w) in enumerate(zip(row.cells, row_data, col_widths)):
            set_cell_bg(cell, fill)
            set_cell_borders(cell, color="CCCCCC", sz=4)
            set_cell_margins(cell)
            set_col_width(cell, w)
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            run.font.color.rgb = DARK_GREY
            run.font.name = "Arial"
            if ci == 0:
                run.bold = True

    doc.add_paragraph()  # spacer


def simple_table(doc, headers, rows_data, col_widths=None):
    total = sum(col_widths) if col_widths else 9360
    table = doc.add_table(rows=1 + len(rows_data), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = 'Table Grid'

    hrow = table.rows[0]
    for i, (cell, hdr) in enumerate(zip(hrow.cells, headers)):
        set_cell_bg(cell, MID_BLUE)
        set_cell_borders(cell, color="FFFFFF", sz=4)
        set_cell_margins(cell)
        if col_widths:
            set_col_width(cell, col_widths[i])
        p = cell.paragraphs[0]
        run = p.add_run(hdr)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(9)
        run.font.name = "Arial"

    for ri, row_data in enumerate(rows_data):
        row = table.rows[ri + 1]
        fill = PALE_BLUE if ri % 2 == 0 else WHITE
        for ci, (cell, val) in enumerate(zip(row.cells, row_data)):
            if col_widths:
                set_col_width(cell, col_widths[ci])
            set_cell_bg(cell, fill)
            set_cell_borders(cell, color="CCCCCC", sz=4)
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            run.font.color.rgb = DARK_GREY
            run.font.name = "Arial"
            if ci == 0:
                run.bold = True

    doc.add_paragraph()


# ════════════════════════════════════════════════════════════════════════════
#  BUILD DOCUMENT
# ════════════════════════════════════════════════════════════════════════════

doc = Document()

# Page size – A4
section = doc.sections[0]
section.page_width  = Cm(21.0)
section.page_height = Cm(29.7)
section.left_margin   = Cm(2.0)
section.right_margin  = Cm(2.0)
section.top_margin    = Cm(2.0)
section.bottom_margin = Cm(2.0)

# ── Cover / Title block ──────────────────────────────────────────────────────
cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_spacing(cover, before=30, after=4)
r = cover.add_run("WORK INSTRUCTION")
r.bold = True; r.font.size = Pt(10); r.font.color.rgb = MID_BLUE; r.font.name = "Arial"

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_spacing(title_p, before=2, after=6)
r = title_p.add_run("Professional 3D Viewer:\nCamera Controls, Navigation & Section Tools")
r.bold = True; r.font.size = Pt(22); r.font.color.rgb = DARK_BLUE; r.font.name = "Arial"

add_para_bottom_border(title_p, color=hex_to_rgb_str(MID_BLUE), size=12)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_spacing(meta, before=6, after=24)
r = meta.add_run("Full specification for implementation in Three.js / WebGL\n"
                 "Target: Navisworks / Revit / AutoCAD-grade interaction quality")
r.font.size = Pt(10); r.font.color.rgb = DARK_GREY; r.font.name = "Arial"

# ── 0. SCOPE ──────────────────────────────────────────────────────────────────
h1(doc, "0  Scope & Objectives")
body(doc,
     "This document specifies every interactive camera, navigation, selection and section control "
     "required to deliver a professional-grade 3D viewer experience. "
     "The AI developer must implement all items below to the stated defaults; every setting must "
     "be exposed in a persistent Settings panel and (where applicable) a toolbar button.")
bullet(doc, "Framework reference: Three.js r160+ / WebGL 2.0")
bullet(doc, "Coordinate convention default: Z-up (CAESAR / AutoCAD) with run-time toggle to Y-up")
bullet(doc, "Input devices: mouse + keyboard (primary), touchscreen (secondary), trackpad")
bullet(doc, "Performance target: ≥60 fps on models up to 500 000 geometry nodes")

# ── 1. CAMERA SYSTEM ─────────────────────────────────────────────────────────
h1(doc, "1  Camera System")

h2(doc, "1.1  Projection Modes")
body(doc, "Both modes must be available and switchable at any time without resetting the view target.")
bullet(doc, "Perspective – depth foreshortening, natural construction view")
bullet(doc, "Orthographic – parallel projection, accurate dimensioning view")
note(doc, "When switching projection, preserve the approximate field-of-view by matching the "
          "orthographic frustum height to the perspective view's visible height at the target distance.")

h3(doc, "Perspective Camera Settings")
settings_table(doc, [
    ("fov",              "Number",  "60",         "Vertical field of view in degrees (10–170)"),
    ("near",             "Number",  "0.1",        "Near clipping plane in world units. Never set to 0."),
    ("far",              "Number",  "1 000 000",  "Far clipping plane in world units"),
    ("aspect",           "Number",  "auto",       "Width / height – updated on every resize event"),
    ("zoom",             "Number",  "1.0",        "Camera zoom multiplier (perspective zoom-in)"),
    ("focalLength",      "Number",  "—",          "Alternative to FOV – 35 mm equivalent focal length (optional exposure)"),
])

h3(doc, "Orthographic Camera Settings")
settings_table(doc, [
    ("left / right",     "Number",  "±span/2",    "Horizontal frustum bounds in world units"),
    ("top / bottom",     "Number",  "±span/2",    "Vertical frustum bounds in world units"),
    ("near",             "Number",  "-100 000",   "Near plane – set negative to see geometry behind camera"),
    ("far",              "Number",  "1 000 000",  "Far clipping plane"),
    ("zoom",             "Number",  "1.0",        "Scale multiplier; increase to zoom in"),
])

note(doc, "Near plane must be validated: near < far at all times; enforce min near = 1e-4 for "
          "perspective and near = -far for ortho to prevent z-fighting artifacts.")

h2(doc, "1.2  Clipping Planes (Near / Far) UI Controls")
body(doc, "Expose the following controls in the Settings panel under Camera > Clipping:")
bullet(doc, "Near Plane slider  –  logarithmic scale, range 0.001 → far−1, step adaptive")
bullet(doc, "Far Plane slider   –  logarithmic scale, range near+1 → 10 000 000, step adaptive")
bullet(doc, "Auto-fit button    –  computes tight near/far from bounding box of visible objects")
bullet(doc, "Reset button       –  restores default values listed above")
note(doc, "Auto-fit near = max(0.01, bbox_min_depth × 0.5); auto-fit far = bbox_max_depth × 2.0. "
          "Recompute after any load or section change.")

h2(doc, "1.3  Standard Orthographic Views")
body(doc, "Toolbar or keyboard shortcut must snap camera to these canonical positions. "
          "Maintain current orbit target; only update camera position and up-vector.")
simple_table(doc,
    ["View Name",       "Camera Direction (world)",     "Up Vector",      "Shortcut"],
    [
        ("Top (Plan)",      "+Z looking down (Z-up)",       "+Y or +X",       "Numpad 7"),
        ("Bottom",          "-Z looking up",                "+Y or +X",       "Ctrl+Numpad 7"),
        ("Front (N.Elev)",  "+Y looking south (Z-up)",      "+Z",             "Numpad 1"),
        ("Back (S.Elev)",   "-Y looking north",             "+Z",             "Ctrl+Numpad 1"),
        ("Right (E.Elev)",  "+X looking west",              "+Z",             "Numpad 3"),
        ("Left (W.Elev)",   "-X looking east",              "+Z",             "Ctrl+Numpad 3"),
        ("ISO NW",          "(-1,+1,+1) normalised",        "+Z",             "Numpad 0"),
        ("ISO NE",          "(+1,+1,+1) normalised",        "+Z",             "Ctrl+Numpad 0"),
        ("Home",            "Fit all visible objects",      "+Z (default)",   "H"),
    ],
    col_widths=[1800, 2800, 1400, 1600])

# ── 2. NAVIGATION MODES ───────────────────────────────────────────────────────
h1(doc, "2  Navigation Modes")
body(doc, "Navigation modes are mutually exclusive. The active mode is indicated by a highlighted "
          "toolbar button. Each mode binds mouse buttons and modifier keys as defined below. "
          "All modes must support pinch-to-zoom and two-finger pan on touch devices.")

h2(doc, "2.1  3D Free Orbit (Default)")
body(doc, "Unconstrained spherical orbit around the current target point.")
simple_table(doc,
    ["Input", "Action"],
    [
        ("Left drag",               "Rotate azimuth (left/right) and elevation (up/down) simultaneously"),
        ("Right drag",              "Pan – translate target and camera in the view plane"),
        ("Scroll wheel",            "Dolly in/out along view axis (perspective) or zoom frustum (ortho)"),
        ("Middle drag",             "Pan (same as right drag)"),
        ("Double-click (L)",        "Set new orbit target at hit point"),
        ("F key",                   "Frame / fit selected objects"),
        ("Shift + scroll",          "Roll camera ±5° per notch"),
    ],
    col_widths=[2800, 6400])
settings_table(doc, [
    ("rotateSpeed",    "Number", "1.0",   "Multiplier for orbit rotation sensitivity"),
    ("panSpeed",       "Number", "1.0",   "Multiplier for pan translation sensitivity"),
    ("zoomSpeed",      "Number", "1.0",   "Multiplier for dolly/zoom speed"),
    ("dampingFactor",  "Number", "0.08",  "Inertia smoothing factor (0 = instant, 1 = no movement)"),
    ("enableDamping",  "Bool",   "true",  "Enable/disable inertia smoothing"),
    ("minDistance",    "Number", "0.01",  "Minimum dolly distance from target (perspective)"),
    ("maxDistance",    "Number", "∞",     "Maximum dolly distance from target"),
    ("invertX",        "Bool",   "false", "Invert left/right orbit direction"),
    ("invertY",        "Bool",   "false", "Invert up/down orbit direction"),
    ("zoomToCursor",   "Bool",   "true",  "Zoom towards mouse cursor position, not orbit target"),
])

h2(doc, "2.2  Planar Rotation (2D Orbit) – About Vertical Axis")
body(doc, "Locks the camera elevation angle; only azimuth (compass direction) changes. "
          "The camera stays at the same height above the target – like spinning a globe. "
          "Used for plan-view navigation.")
bullet(doc, "Left drag left/right  →  azimuth changes; elevation is frozen at the angle when mode was activated")
bullet(doc, "Elevation changes are ignored (Y-axis / vertical-axis rotation only)")
bullet(doc, "Right drag and scroll behave identically to 3D Orbit mode")
bullet(doc, "Implementation: bypass OrbitControls rotation – apply applyAxisAngle(VERTICAL_AXIS, angle) "
            "directly on camera offset vector. Do NOT use minPolarAngle = maxPolarAngle constraint "
            "(unreliable in OrbitControls).")
note(doc, "The "Rotate about X-axis" (plan) variant first snaps to the top-down view at 18° "
          "elevation, then activates Planar Rotation. The "Rotate about Y-axis" variant activates "
          "Planar Rotation from the current view without snapping.")

h2(doc, "2.3  Axis-Constrained Rotation")
body(doc, "Three dedicated modes that lock rotation to exactly one world axis. "
          "Toolbar provides a button for each.")
simple_table(doc,
    ["Mode",                   "Locked Axis",  "Free Movement",                  "Visual Cue"],
    [
        ("Rotate about X axis", "X (East)",    "Pitch (nod) around East axis",    "Red X highlight in gizmo"),
        ("Rotate about Y axis", "Y (North)",   "Yaw (turn) around North axis",    "Green Y highlight"),
        ("Rotate about Z axis", "Z (Up/Down)", "Roll around vertical axis",       "Blue Z highlight"),
        ("Planar (free azimuth)","Z (vertical)","Azimuth only – elevation frozen", "ViewCube compass ring"),
    ],
    col_widths=[2200, 1500, 3200, 2300])
note(doc, "For each constrained mode, compute the rotation delta from horizontal or vertical mouse "
          "delta only (whichever is primary), then rotate camera.position – controls.target about "
          "the chosen world axis using applyAxisAngle. Never use OrbitControls polar/azimuth clamp "
          "for these – implement direct vector math instead.")

h2(doc, "2.4  Pan")
body(doc, "Translates the camera and orbit target together in the view plane (no rotation). "
          "Can be a dedicated mode or always available via right-drag in other modes.")
bullet(doc, "Speed scales with distance to target so pan feels uniform at any zoom level")
bullet(doc, "Hold Shift in any orbit mode to temporarily activate pan")
settings_table(doc, [
    ("panSpeed",        "Number", "1.0",   "Pan sensitivity multiplier"),
    ("screenSpacePanning","Bool", "true",  "true = pan in screen plane; false = pan in XZ world plane"),
    ("panBounds",       "AABB",  "none",   "Optional: constrain pan to keep target inside a world bounding box"),
])

h2(doc, "2.5  Zoom / Dolly")
body(doc, "Two zoom strategies depending on projection:")
bullet(doc, "Perspective: move camera closer/farther along view axis (dolly). Near/far auto-adjust.")
bullet(doc, "Orthographic: scale left/right/top/bottom of frustum symmetrically about cursor point.")
settings_table(doc, [
    ("zoomSpeed",       "Number", "1.0",   "Speed multiplier for scroll wheel zoom"),
    ("minZoom",         "Number", "1e-4",  "Minimum ortho zoom value (prevents collapse)"),
    ("maxZoom",         "Number", "∞",     "Maximum ortho zoom value"),
    ("zoomToCursor",    "Bool",   "true",  "Recompute orbit target at cursor's world hit point on zoom"),
    ("autoNearFar",     "Bool",   "true",  "Auto-adjust near/far planes during dolly to prevent clipping"),
])

h2(doc, "2.6  Fly / Walk Mode (Optional – Advanced)")
body(doc, "First-person free-fly for large models. Toggle with F9 or toolbar button.")
bullet(doc, "WASD / arrow keys – move forward, back, strafe left, strafe right")
bullet(doc, "E / Q – move up / down")
bullet(doc, "Mouse move (left button held) – look direction")
bullet(doc, "Scroll wheel – change movement speed")
bullet(doc, "Shift – sprint (5× speed)")
settings_table(doc, [
    ("flySpeed",      "Number", "100",   "Base movement speed in world units per second"),
    ("lookSensitivity","Number","0.15",  "Mouse look sensitivity degrees per pixel"),
    ("enableCollision","Bool",  "false", "Prevent camera from passing through geometry"),
])

# ── 3. SELECTION SYSTEM ───────────────────────────────────────────────────────
h1(doc, "3  Selection System")

h2(doc, "3.1  Click-to-Select")
body(doc, "Left-click on any geometry in Select mode picks the object via raycasting.")
bullet(doc, "Highlight picked object with emissive colour overlay (default: #FFA500 orange, 50% opacity)")
bullet(doc, "Display object name, type, tag and bounding box dimensions in a side panel")
bullet(doc, "Click on empty space deselects all")
bullet(doc, "Ctrl + click  – add/remove from selection set (multi-select)")
bullet(doc, "Shift + click  – range select (all objects between last pick and this pick)")
settings_table(doc, [
    ("highlightColor",    "HEX",    "#FFA500", "Emissive highlight colour for selected objects"),
    ("hoverColor",        "HEX",    "#88CCFF", "Highlight colour on hover (before click)"),
    ("selectionOpacity",  "Number", "0.5",     "Blending opacity of selection overlay"),
    ("enableHover",       "Bool",   "true",    "Show hover highlight as mouse moves"),
    ("selectThruTransparent","Bool","false",   "Whether raycaster picks through transparent surfaces"),
])

h2(doc, "3.2  Box Select")
body(doc, "Hold Alt + left-drag to draw a screen-space rectangle. "
          "All objects whose projected bounding box intersects the rectangle are selected.")
bullet(doc, "Rectangle drawn as a dashed overlay on the canvas (2 px dashed line, MID_BLUE colour)")
bullet(doc, "Release mouse to commit selection")
bullet(doc, "Hold Alt + Ctrl to ADD to existing selection")

h2(doc, "3.3  Rotate About Selection")
body(doc, "When objects are selected, pressing shortcut key 3 (or toolbar button) sets the orbit "
          "pivot to the centroid of the selection's bounding box instead of the model centre.")
bullet(doc, "Visual indicator: small 3-axis cross-hair gizmo drawn at pivot point")
bullet(doc, "Orbit, Planar Rotation and all constrained modes use this pivot until reset")
bullet(doc, "Reset pivot  –  shortcut R or button; returns target to model bounding-box centre")
bullet(doc, "Double-click on any surface also sets the pivot to the hit point (surface snapping)")
note(doc, "Store pivot in controls.target. All navigation modes must read this value. "
          "Do not hard-code model centre as the orbit origin.")

h2(doc, "3.4  Select & Isolate")
bullet(doc, "I key or context menu – hide all non-selected objects; fit camera to selection")
bullet(doc, "Esc or button – restore all hidden objects")

# ── 4. AXIS / COORDINATE SYSTEM ───────────────────────────────────────────────
h1(doc, "4  World Axis & Coordinate System")

h2(doc, "4.1  Vertical Axis Toggle")
body(doc, "The application must support both conventions. The setting is stored in localStorage "
          "and applied on startup.")
simple_table(doc,
    ["Convention", "Vertical Axis", "Typical Use",         "Three.js camera.up"],
    [
        ("Z-up (default)", "Z (blue)",  "CAESAR, AutoCAD, Civil3D, Revit", "new Vector3(0,0,1) + custom quaternion"),
        ("Y-up",           "Y (green)", "Three.js native, glTF, Maya",     "new Vector3(0,1,0)"),
    ],
    col_widths=[1800, 1400, 3200, 2800])
body(doc, "When switching axis convention:")
bullet(doc, "Rotate all geometry 90° about world X axis (or apply a scene root quaternion – do not modify per-mesh transforms)")
bullet(doc, "Update camera.up vector")
bullet(doc, "Update OrbitControls orbit axis")
bullet(doc, "Update axis gizmo colours and labels")
bullet(doc, "Recompute standard view presets (Top = looking along new −vertical)")

h2(doc, "4.2  Axis Gizmo (World Indicator)")
body(doc, "A fixed-size, non-interactive 3-axis indicator in the bottom-left viewport corner "
          "(or as chosen in Settings). It rotates synchronously with the camera.")
settings_table(doc, [
    ("showAxisGizmo",   "Bool",  "true",           "Show / hide the axis indicator"),
    ("gizmoSize",       "Number","80 px",           "Pixel size of the gizmo widget"),
    ("gizmoPosition",   "Enum",  "bottom-left",     "Corner: top-left, top-right, bottom-left, bottom-right"),
    ("labelAxes",       "Bool",  "true",            "Show X / Y / Z labels on each axis"),
    ("xColor",          "HEX",   "#E74C3C (red)",   "Colour of X axis line and label"),
    ("yColor",          "HEX",   "#2ECC71 (green)", "Colour of Y axis line and label"),
    ("zColor",          "HEX",   "#3498DB (blue)",  "Colour of Z axis line and label"),
    ("gizmoClickable",  "Bool",  "true",            "Click axis to snap camera to that standard view"),
])

h2(doc, "4.3  Ground Grid")
settings_table(doc, [
    ("showGrid",         "Bool",   "true",    "Toggle ground-plane grid visibility"),
    ("gridSize",         "Number", "10 000",  "Total grid extent in world units"),
    ("gridDivisions",    "Number", "100",     "Number of grid cells per axis"),
    ("gridColor",        "HEX",    "#888888", "Grid line colour"),
    ("gridCenterColor",  "HEX",    "#444444", "Centre cross-hair colour"),
    ("gridFade",         "Bool",   "true",    "Fade grid lines at distance (alpha fall-off)"),
    ("gridPlane",        "Enum",   "XY (Z-up)","Which world plane the grid lies on"),
])

# ── 5. VIEW CUBE ──────────────────────────────────────────────────────────────
h1(doc, "5  ViewCube")
body(doc, "The ViewCube is an interactive 3D cube widget (top-right corner by default) that shows "
          "the camera orientation at all times and allows single-click navigation to any standard view.")

h2(doc, "5.1  Interaction Model")
simple_table(doc,
    ["Click Target",      "Action"],
    [
        ("Face (e.g. FRONT)", "Snap camera to the orthographic view looking at that face (smooth animated transition)"),
        ("Edge (12 total)",   "Snap to 45° isometric view along that edge"),
        ("Corner (8 total)",  "Snap to trimetric view along that corner diagonal"),
        ("Home icon",         "Fit-all: frame entire model in current projection"),
        ("Compass ring",      "Click and drag to rotate about the vertical axis only (planar rotation)"),
        ("Compass arrow",     "Click N/S/E/W label to snap to that elevation view"),
        ("Cube drag",         "Click and drag the cube body to freely orbit (same as 3D Orbit mode)"),
    ],
    col_widths=[2400, 6800])

h2(doc, "5.2  ViewCube Settings")
settings_table(doc, [
    ("showViewCube",         "Bool",   "true",          "Show / hide the ViewCube"),
    ("viewCubeSize",         "Number", "120 px",        "Pixel size of cube widget"),
    ("viewCubePosition",     "Enum",   "top-right",     "Corner: top-left, top-right, bottom-left, bottom-right"),
    ("viewCubeOpacity",      "Number", "0.85",          "Default opacity (1.0 on hover)"),
    ("viewCubeAnimDuration", "Number", "400 ms",        "Snap animation duration in milliseconds"),
    ("viewCubeAnimEasing",   "Enum",   "ease-in-out",   "Easing function for snap animation"),
    ("showCompass",          "Bool",   "true",          "Show compass rose below the cube"),
    ("northAngle",           "Number", "0°",            "Rotation of North label relative to model X+ axis"),
    ("viewCubeFaceLabels",   "String", "TOP,BOTTOM,etc","Comma-separated face labels (localisation-ready)"),
    ("viewCubeHighlightHover","Bool",  "true",          "Highlight face/edge/corner under mouse cursor"),
])

# ── 6. NAVIGATION TOOLBAR ─────────────────────────────────────────────────────
h1(doc, "6  Navigation Toolbar")
body(doc, "A vertical (or horizontal) floating toolbar with icon buttons for each mode and view preset. "
          "Active mode button is highlighted. Tooltip on hover. All icons must be SVG for crisp "
          "rendering at all DPI settings.")

h2(doc, "6.1  Required Toolbar Buttons")
simple_table(doc,
    ["Button",               "Mode / Action",              "Shortcut", "Icon Description"],
    [
        ("Select",           "Select mode",                "S",        "Arrow cursor with box"),
        ("3D Orbit",         "Free 3D orbit",              "O",        "Sphere with orbit arc"),
        ("Rotate about X",   "Planar rotation + snap Top", "X",        "Flat circle / plan arc"),
        ("Rotate about Y",   "Planar rot. current view",   "Y",        "Vertical orbit arc"),
        ("Pan",              "Pan mode",                   "P",        "Four-way arrow"),
        ("Zoom Window",      "Drag zoom rectangle",        "Z",        "Magnifier + rectangle"),
        ("Fit All",          "Frame entire model",         "H",        "Fit-to-screen brackets"),
        ("Fit Selection",    "Frame selected objects",     "F",        "Target with selection"),
        ("Top View",         "Snap top orthographic",      "Numpad 7", "Top arrow"),
        ("Front View",       "Snap front elevation",       "Numpad 1", "Front arrow"),
        ("Right View",       "Snap right elevation",       "Numpad 3", "Right arrow"),
        ("ISO View",         "Snap isometric NW",          "Numpad 0", "Cube corner"),
        ("Perspective/Ortho","Toggle projection",          "V",        "Perspective lines vs parallel"),
        ("Section Box",      "Toggle section box",         "B",        "Box with dashed cut"),
        ("Settings",         "Open settings panel",        ",",        "Gear icon"),
    ],
    col_widths=[1900, 2200, 1400, 3700])

h2(doc, "6.2  Toolbar Settings")
settings_table(doc, [
    ("toolbarPosition",   "Enum",   "left",         "left | right | top | bottom"),
    ("toolbarIconSize",   "Number", "32 px",        "Icon button size in pixels"),
    ("toolbarOpacity",    "Number", "0.88",         "Background opacity of toolbar panel"),
    ("showTooltips",      "Bool",   "true",         "Show tooltip labels on hover"),
    ("tooltipDelay",      "Number", "500 ms",       "Milliseconds before tooltip appears"),
])

# ── 7. SECTION / CLIPPING ──────────────────────────────────────────────────────
h1(doc, "7  Section & Clipping Planes")

h2(doc, "7.1  Near / Far Clipping (see also §1.2)")
body(doc, "Already covered in §1.2. Surfaces within the near-far frustum are visible; "
          "everything outside is clipped by the GPU. Expose sliders in the Settings panel.")

h2(doc, "7.2  Six-Plane Section Box")
body(doc, "A box-shaped clipping region defined by six world-axis-aligned planes (±X, ±Y, ±Z). "
          "Geometry outside the box is clipped in real-time using Three.js clipping planes.")
bullet(doc, "Toggle entire section box on/off with toolbar button or shortcut B")
bullet(doc, "Each face of the box is independently draggable via 3D handles")
bullet(doc, "Visual: translucent box outline + coloured handles on each face")
bullet(doc, "Flip any plane with a double-click to clip inside-out")
bullet(doc, "Section cap: fill the cut cross-section with a solid colour")
settings_table(doc, [
    ("enableSectionBox",   "Bool",   "false",          "Global toggle for the section box"),
    ("sectionBoxPadding",  "Number", "5% of bbox",     "Initial padding from model bounding box when first enabled"),
    ("clipX_min",          "Number", "−∞",             "Left clipping plane position (world X)"),
    ("clipX_max",          "Number", "+∞",             "Right clipping plane position (world X)"),
    ("clipY_min",          "Number", "−∞",             "Front clipping plane position (world Y)"),
    ("clipY_max",          "Number", "+∞",             "Back clipping plane position (world Y)"),
    ("clipZ_min",          "Number", "−∞",             "Bottom clipping plane position (world Z)"),
    ("clipZ_max",          "Number", "+∞",             "Top clipping plane position (world Z)"),
    ("showSectionCap",     "Bool",   "true",           "Fill cut cross-section with cap colour"),
    ("sectionCapColor",    "HEX",    "#E0E0E0",        "Section cap fill colour"),
    ("sectionCapOpacity",  "Number", "0.7",            "Section cap opacity"),
    ("sectionBoxColor",    "HEX",    "#2E75B6",        "Section box outline colour"),
    ("sectionBoxOpacity",  "Number", "0.15",           "Section box face transparency"),
    ("sectionHandleSize",  "Number", "12 px",          "Pixel size of draggable handles on box faces"),
])

h2(doc, "7.3  Single Clipping Planes")
body(doc, "In addition to the section box, support up to 6 individually configurable clipping planes "
          "at arbitrary orientations (not necessarily axis-aligned).")
settings_table(doc, [
    ("plane_N_normal",   "Vector3","varies",    "Unit normal of Nth clipping plane"),
    ("plane_N_constant", "Number", "0",         "Plane constant d in ax+by+cz+d=0"),
    ("plane_N_enabled",  "Bool",   "false",     "Enable/disable Nth plane independently"),
    ("plane_N_flip",     "Bool",   "false",     "Invert clipping direction for this plane"),
    ("clipIntersection", "Bool",   "false",     "false = union clip; true = intersection clip"),
])

# ── 8. SETTINGS PANEL ──────────────────────────────────────────────────────────
h1(doc, "8  Settings Panel")
body(doc, "A collapsible side panel (or modal) exposing all configurable parameters. "
          "Settings are auto-saved to localStorage under key 'viewer3d_settings'. "
          "A Reset All button restores all defaults. Export/Import JSON buttons allow "
          "sharing settings between sessions.")

h2(doc, "8.1  Panel Sections")
bullet(doc, "Camera  –  near, far, FOV, projection, zoom-to-cursor, auto near/far")
bullet(doc, "Navigation  –  rotate speed, pan speed, zoom speed, damping, invert axes, min/max distance")
bullet(doc, "Appearance  –  background colour, grid, axis gizmo, ViewCube, shadows")
bullet(doc, "Selection  –  highlight colour, hover colour, opacity")
bullet(doc, "Section  –  all section box and clipping plane settings (§7)")
bullet(doc, "Performance  –  LOD threshold, shadow resolution, antialias, pixel ratio")
bullet(doc, "Coordinate System  –  vertical axis (Y-up / Z-up), north angle")
bullet(doc, "Keyboard Shortcuts  –  read-only reference table of all shortcuts")

h2(doc, "8.2  Global Settings Reference")
settings_table(doc, [
    ("theme",                "Enum",   "dark",           "UI theme: dark | light | system"),
    ("backgroundColor",      "HEX",   "#1A1A2E",         "Viewport background colour"),
    ("antialias",            "Bool",  "true",            "Enable WebGL MSAA antialiasing"),
    ("shadowsEnabled",       "Bool",  "false",           "Enable real-time shadows (performance cost)"),
    ("pixelRatio",           "Number","window.dpr",      "Renderer pixel ratio (1 for performance, 2 for HiDPI)"),
    ("verticalAxis",         "Enum",  "Z",               "World vertical axis: Y or Z"),
    ("coordinateLabels",     "Bool",  "true",            "Show X/Y/Z or E/N/Up labels on gizmo"),
    ("persistentSettings",   "Bool",  "true",            "Auto-save settings to localStorage"),
])

# ── 9. KEYBOARD SHORTCUTS ──────────────────────────────────────────────────────
h1(doc, "9  Keyboard Shortcuts")
body(doc, "All shortcuts are re-bindable via Settings > Keyboard. Shown here are defaults.")
simple_table(doc,
    ["Shortcut",           "Action"],
    [
        ("O",              "Activate 3D Orbit mode"),
        ("S",              "Activate Select mode"),
        ("X",              "Rotate about X axis (plan snap)"),
        ("Y",              "Rotate about Y axis (current view)"),
        ("P",              "Pan mode"),
        ("Z",              "Zoom window"),
        ("H",              "Fit all (home)"),
        ("F",              "Fit selected"),
        ("V",              "Toggle Perspective / Orthographic"),
        ("B",              "Toggle Section Box"),
        ("I",              "Isolate selected"),
        ("Esc",            "Deselect all / exit Isolate / cancel operation"),
        ("3",              "Set orbit pivot to selection centroid"),
        ("R",              "Reset orbit pivot to model centre"),
        ("G",              "Toggle ground grid"),
        ("A",              "Toggle axis gizmo"),
        ("Numpad 7",       "Top view"),
        ("Numpad 1",       "Front view"),
        ("Numpad 3",       "Right view"),
        ("Numpad 0",       "ISO view"),
        ("Ctrl+Numpad 7",  "Bottom view"),
        ("Ctrl+Numpad 1",  "Back view"),
        ("Ctrl+Numpad 3",  "Left view"),
        ("Ctrl+Z",         "Undo last view change"),
        ("Ctrl+Y",         "Redo view change"),
        ("F9",             "Toggle Fly / Walk mode"),
        (",",              "Open / close Settings panel"),
    ],
    col_widths=[2200, 7000])

# ── 10. IMPLEMENTATION NOTES ──────────────────────────────────────────────────
h1(doc, "10  Implementation Notes (Three.js Specific)")

h2(doc, "10.1  Constrained Rotation – Do NOT Use OrbitControls Clamping")
body(doc, "Three.js OrbitControls clamps (minPolarAngle = maxPolarAngle or "
          "minAzimuthAngle = maxAzimuthAngle) are unreliable when set to equal values – "
          "they can freeze all rotation. Use direct vector math instead:")
bullet(doc, "Set controls.enableRotate = false to suppress OrbitControls rotation processing")
bullet(doc, "Register pointer events (pointerdown / pointermove / pointerup) on the renderer canvas")
bullet(doc, "On pointermove: compute delta, call offset.applyAxisAngle(AXIS, angle), "
            "then camera.position = controls.target + offset; camera.lookAt(controls.target)")
bullet(doc, "On mode exit: remove listeners; restore controls.enableRotate = true")
bullet(doc, "Keep controls.enabled = true throughout so pan (right drag) and zoom (scroll) "
            "still work via OrbitControls unmodified")

h2(doc, "10.2  Z-up Coordinate System")
body(doc, "Three.js is Y-up by default. To support Z-up without modifying every mesh:")
bullet(doc, "Create a scene root Group; rotate it: sceneRoot.rotation.x = -Math.PI / 2")
bullet(doc, "Set camera.up = new THREE.Vector3(0, 0, 1)")
bullet(doc, "Wrap OrbitControls: override orbit axis from Y to Z in the update function or "
            "subclass OrbitControls and change the up reference vector")
bullet(doc, "Apply the same rotation to the orbit axis gizmo and ViewCube")

h2(doc, "10.3  Near / Far Auto-Fit Algorithm")
body(doc, "Call this function after every load, every section change, and after any camera move "
          "that changes the bounding sphere visibility:")
body(doc,
     "1. Collect all visible mesh world positions → compute scene bounding sphere (centre C, radius R)\n"
     "2. dist = camera.position.distanceTo(C)\n"
     "3. near = max(NEAR_MIN, dist − R) × 0.1\n"
     "4. far  = (dist + R) × 10\n"
     "5. camera.near = near; camera.far = far; camera.updateProjectionMatrix()",
     indent=True)

h2(doc, "10.4  Section Cap Rendering")
body(doc, "Three.js does not render section caps natively. Use the stencil buffer technique:")
bullet(doc, "Render scene normally with clipping planes active")
bullet(doc, "Render a full-screen quad masked by stencil to fill cut surfaces with cap material")
bullet(doc, "Alternatively, use a BackSide mesh scaled slightly larger as a simple approximation")

h2(doc, "10.5  Performance Checklist")
bullet(doc, "Merge static geometry with BufferGeometryUtils.mergeGeometries()")
bullet(doc, "Use Level of Detail (LOD) objects for elements > 1 000 mm from camera")
bullet(doc, "Frustum-cull objects before adding to scene (Three.js does this automatically per frame)")
bullet(doc, "Throttle raycasting for hover selection to every other frame (requestAnimationFrame odd/even)")
bullet(doc, "Use WebGLRenderTarget for off-screen picking (GPU color-picking) when object count > 10 000")

# ── 11. ACCEPTANCE CRITERIA ───────────────────────────────────────────────────
h1(doc, "11  Acceptance Criteria")
body(doc, "The implementation is complete when ALL of the following pass:")
simple_table(doc,
    ["#",  "Test",                                                        "Pass condition"],
    [
        ("1",  "3D Orbit from any view",              "Camera rotates smoothly in all directions with damping"),
        ("2",  "Planar rotation from N.Elev",         "Azimuth changes; elevation stays constant at ±0.5°"),
        ("3",  "Rotate-about-Y from iso view",        "Same as #2; no snap to top view"),
        ("4",  "Click to select",                     "Object highlighted; info panel populated"),
        ("5",  "Rotate about selection (key 3)",      "Orbit pivots at selection centroid, not model centre"),
        ("6",  "ViewCube face click",                 "Camera animates to correct orthographic view in ≤ 500 ms"),
        ("7",  "Near/far auto-fit",                   "No z-fighting or near-clipping on any standard view"),
        ("8",  "Section box drag",                    "Geometry clips in real time at 60 fps"),
        ("9",  "Section cap visible",                 "Cut cross-section filled with cap colour"),
        ("10", "Z-up / Y-up toggle",                  "All geometry, gizmo, ViewCube and grid update without reload"),
        ("11", "Settings persist",                    "Reload page – all settings restored from localStorage"),
        ("12", "Touch orbit",                         "Single finger drags orbit; two-finger pinch zooms"),
        ("13", "Keyboard shortcuts",                  "All shortcuts in §9 produce correct actions"),
        ("14", "60 fps performance",                  "Model with 200 k meshes maintains ≥ 60 fps on orbit"),
    ],
    col_widths=[300, 3500, 5400])

# ── Footer note ───────────────────────────────────────────────────────────────
doc.add_paragraph()
hr = doc.add_paragraph()
add_para_bottom_border(hr, color=hex_to_rgb_str(MID_BLUE), size=4)
set_spacing(hr, after=4)

footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = footer_p.add_run("This document is a living specification – update version date when any section changes. "
                      "All numeric defaults may be tuned during QA without requiring document revision.")
r.font.size = Pt(8)
r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
r.font.name = "Arial"

# ── Save ──────────────────────────────────────────────────────────────────────
out = r"C:\Code\CRF-3\3D_Viewer_Camera_Navigation_Work_Instruction.docx"
doc.save(out)
print(f"Saved: {out}")
